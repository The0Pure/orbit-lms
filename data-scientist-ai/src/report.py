"""Generates a professional, business-oriented Word (.docx) report."""
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x29, 0x3B)
GOLD = RGBColor(0xB8, 0x96, 0x5A)
GREY = RGBColor(0x6B, 0x72, 0x80)


def _label(metric: str) -> str:
    return metric.replace("_", " ").title()


def _heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _shade_cell(hdr_cells[i], "1E293B")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def _trend_word(pct: float) -> str:
    if pct >= 15:
        return "strong growth"
    if pct >= 2:
        return "modest growth"
    if pct > -2:
        return "a flat trajectory"
    if pct > -15:
        return "a modest decline"
    return "a significant decline"


def _recommendation(metric_label: str, pct: float) -> str:
    if pct >= 15:
        return (
            f"{metric_label} is on a strong upward trajectory. This is a good time to invest "
            f"further in whatever is driving this metric — consider scaling the underlying "
            f"activity (marketing, capacity, staffing) to capture the momentum before it plateaus."
        )
    if pct >= 2:
        return (
            f"{metric_label} is trending modestly upward. The current approach is working — "
            f"keep monitoring it monthly to confirm the trend holds before committing significant "
            f"new investment."
        )
    if pct > -2:
        return (
            f"{metric_label} is essentially flat. If growth is a priority here, this is a signal "
            f"to test new initiatives rather than expecting the current trajectory to improve on its own."
        )
    if pct > -15:
        return (
            f"{metric_label} shows early signs of decline. It's worth investigating the root cause "
            f"now — a small course correction tends to be cheaper than waiting for the trend to worsen."
        )
    return (
        f"{metric_label} is declining significantly. This warrants immediate attention from "
        f"decision-makers — review what changed around the point the decline started and prioritize "
        f"a corrective action plan."
    )


def build_report(
    df,
    log: dict,
    forecasts: dict,
    yearly: dict,
    chart_titles: list[str],
    dashboard_image: str,
    output_path: str,
):
    doc = Document()

    # ── Cover ──────────────────────────────────────────────
    title = doc.add_heading("Business Intelligence Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph("Prepared by Data Scientist AI")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GOLD
    sub.runs[0].font.bold = True

    date_p = doc.add_paragraph(f"{date.today().strftime('%B %d, %Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = GREY

    doc.add_paragraph()

    # ── Executive Summary ──────────────────────────────────
    _heading(doc, "Executive Summary")
    retained_pct = (log["rows_out"] / log["rows_in"] * 100) if log["rows_in"] else 100
    quality_note = (
        "the dataset was already clean" if log["duplicates_removed"] == 0 and not log["missing_filled"]
        else "a number of data quality issues were found and corrected"
    )
    summary_lines = [
        f"This report analyzes {log['rows_in']:,} records. After cleaning, "
        f"{log['rows_out']:,} records ({retained_pct:.0f}%) were retained for analysis; "
        f"{quality_note}."
    ]
    if forecasts:
        metric_names = ", ".join(_label(m) for m in forecasts)
        summary_lines.append(f"Forward-looking projections were generated for: {metric_names}.")
    if yearly:
        growing = [m for m, c in yearly.items() if c["pct_change"] >= 2]
        declining = [m for m, c in yearly.items() if c["pct_change"] <= -2]
        if growing:
            summary_lines.append(
                f"{', '.join(_label(m) for m in growing)} {'is' if len(growing)==1 else 'are'} "
                f"projected to grow next year."
            )
        if declining:
            summary_lines.append(
                f"{', '.join(_label(m) for m in declining)} {'is' if len(declining)==1 else 'are'} "
                f"projected to decline next year and should be reviewed."
            )
    doc.add_paragraph(" ".join(summary_lines))

    # ── Data Overview ──────────────────────────────────────
    _heading(doc, "1. Data Overview")
    doc.add_paragraph(
        "The table below summarizes the scope and quality of the data behind this report."
    )
    overview_rows = [
        ("Records analyzed (before cleaning)", f"{log['rows_in']:,}"),
        ("Records used (after cleaning)", f"{log['rows_out']:,}"),
        ("Duplicate records removed", f"{log['duplicates_removed']:,}"),
        ("Columns with missing values fixed", f"{len(log['missing_filled'])}"),
        ("Date column used for trends", log.get("date_column") or "None detected"),
    ]
    _add_table(doc, ["Metric", "Value"], overview_rows)

    if log["missing_filled"]:
        doc.add_paragraph()
        doc.add_paragraph(
            "Missing data was filled using sensible defaults (the typical value for numbers, "
            "\"Unknown\" for text) so that no records had to be discarded:"
        )
        for col, count in log["missing_filled"].items():
            doc.add_paragraph(f"{col.replace('_', ' ').title()}: {count} values filled", style="List Bullet")

    # ── Dashboard ──────────────────────────────────────────
    _heading(doc, "2. Dashboard")
    doc.add_paragraph(
        "The dashboard below gives a visual summary of how each key metric has performed and "
        "where it is headed. Here's what each chart means in plain terms:"
    )
    doc.add_picture(dashboard_image, width=Inches(6.3))

    for metric in forecasts:
        label = _label(metric)
        doc.add_paragraph(
            f"{label} — Trend & Forecast: the solid blue line is what actually happened; the dashed "
            f"orange line is where {label.lower()} is headed next if the recent trend continues. "
            f"Think of the dashed segment as a short-term planning estimate, not a guarantee.",
            style="List Bullet",
        )
    for metric, comp in yearly.items():
        label = _label(metric)
        last_year = comp["years"][-1]["year"]
        direction = "ahead of" if comp["years"][-1]["pct_change"] >= 0 else "behind"
        years_word = "year" if len(comp["years"]) == 1 else f"{len(comp['years'])} years"
        doc.add_paragraph(
            f"{label} — {comp['this_year']} vs {last_year}: side-by-side bars compare this year's "
            f"actual monthly totals (blue) against each of the next {years_word}' projected totals "
            f"(orange shades), so you can see at a glance which months are expected to run "
            f"{direction} this year's pace.",
            style="List Bullet",
        )
    for t in chart_titles[len(forecasts) + len(yearly):]:
        col = t.replace("Distribution of ", "")
        doc.add_paragraph(
            f"{col} — Distribution: this histogram shows how {col.lower()} values are spread across "
            f"all records — where most of the activity is concentrated, and whether any unusual "
            f"outliers stand out that are worth a closer look.",
            style="List Bullet",
        )

    # ── Performance & Forecast ──────────────────────────────
    _heading(doc, "3. Performance & Forecast")
    if forecasts:
        for metric, payload in forecasts.items():
            label = _label(metric)
            data = payload["data"]
            forecast_rows = data[data["type"] == "forecast"]
            actual_rows = data[data["type"] == "actual"]["value"]
            recent_actual = actual_rows.tail(min(7, len(actual_rows))).mean()
            avg_forecast = forecast_rows["value"].mean()
            change_pct = ((avg_forecast - recent_actual) / recent_actual * 100) if recent_actual else 0
            trend = _trend_word(change_pct)
            doc.add_paragraph(
                f"{label}: based on recent history, {label.lower()} is showing {trend} over the "
                f"next {len(forecast_rows)} days, moving from a recent average level of "
                f"{recent_actual:,.1f} toward a projected average of {avg_forecast:,.1f}."
            )
    else:
        doc.add_paragraph(
            "No numeric metrics with a usable date column were available to forecast in this dataset."
        )

    # ── Year-over-Year Outlook ──────────────────────────────
    _heading(doc, "4. Year-over-Year Outlook")
    if yearly:
        max_years = max(len(comp["years"]) for comp in yearly.values())
        doc.add_paragraph(
            "The table below compares this year's totals against the projection for each "
            f"of the next {max_years} year(s), for each key metric, assuming current trends continue."
        )
        headers = ["Metric", "This Year"] + [
            f"+{i} Year (Forecast)" if i > 1 else "Next Year (Forecast)" for i in range(1, max_years + 1)
        ]
        yoy_rows = []
        for metric, comp in yearly.items():
            row = [_label(metric), f"{comp['this_year_total']:,.1f}"]
            for i in range(max_years):
                if i < len(comp["years"]):
                    y = comp["years"][i]
                    row.append(f"{y['total']:,.1f} ({y['pct_change']:+.1f}%)")
                else:
                    row.append("—")
            yoy_rows.append(tuple(row))
        _add_table(doc, headers, yoy_rows)
    else:
        doc.add_paragraph(
            "Not enough date-stamped history was available to compare this year against a "
            "future-year projection."
        )

    # ── Business Recommendations ─────────────────────────────
    _heading(doc, "5. Business Recommendations")
    if yearly:
        for metric, comp in yearly.items():
            doc.add_paragraph(_recommendation(_label(metric), comp["pct_change"]), style="List Bullet")
    else:
        doc.add_paragraph(
            "Add a date column and at least one numeric metric to your data to receive "
            "tailored, trend-based recommendations in future reports."
        )

    # ── Conclusion ───────────────────────────────────────────
    _heading(doc, "6. Conclusion")
    doc.add_paragraph(
        f"In summary, the data was cleaned to {retained_pct:.0f}% confidence in record quality, "
        + (
            f"trends were modeled for {len(forecasts)} key metric(s), and the year-over-year outlook "
            f"points to actionable opportunities and risks outlined above. Decision-makers should "
            f"revisit this report periodically as new data becomes available to keep the forecasts current."
            if forecasts
            else "though no time-series metrics were available for forecasting in this particular dataset."
        )
    )

    doc.save(output_path)
